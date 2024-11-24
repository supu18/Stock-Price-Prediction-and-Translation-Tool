"""
            Author: Supriya Jayarj
            Date: 2024-02-10

            This script is the main entry point for the file translation project. The tool is a Python-based utility that allows users to translate text content from various file formats. It utilizes Google Translate API for language translation and supports popular file types such as text files (.txt), Word documents (.docx), Excel spreadsheets (.xlsx), CSV files (.csv). The tool also supports the detection of the source language and provides a list of supported languages for translation.
"""
import os
import tkinter as tk
from tkinter import filedialog, messagebox
from shutil import copyfile
import logging
import csv
from langdetect import detect
from googletrans import Translator, LANGUAGES
from docx import Document
from openpyxl import load_workbook, Workbook
# from reportlab.lib.pagesizes import letter
# from reportlab.pdfgen import canvas
# from io import BytesIO
# from PyPDF2.generic import TextStringObject, NameObject, FloatObject, ArrayObject
# from PyPDF2 import PdfReader, PdfWriter, Transformation
# Set up logging
logging.basicConfig(filename='translation.log', level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')


def detect_language(text):
    """
    Detect the language of the given text using langdetect library.

    Parameters:
        text (str): The text to detect language from.

    Returns:
        str: The detected language code.
    """
    try:
        return detect(text)
    except Exception as e:
        logging.error(f"Error occurred while detecting language: {e}")
        messagebox.showerror("Error", "An error occurred while detecting the language.")
        return None


def translate(text, dest_lang):
    """
    Translate the given text to the specified destination language using Google Translate API.

    Parameters:
        text (str): The text to be translated.
        dest_lang (str): The destination language code.

    Returns:
        str: The translated text.
    """
    try:
        translator = Translator(service_urls=[
            'translate.google.com',
            'translate.google.co.kr',
        ])
        translated = translator.translate(text, dest=dest_lang)
        return translated.text
    except Exception as e:
        logging.error(f"Error occurred while translating text: {e}")
        messagebox.showerror("Error", "An error occurred while translating the text.")
        return None


def read_file(file_path):
    """
    Read the contents of a file.

    Parameters:
        file_path (str): The path to the file to be read.

    Returns:
        str: The content of the file.
    """
    _, file_ext = os.path.splitext(file_path)
    text = ""

    try:
        if file_ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
        elif file_ext == '.docx':
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + '\n'
            # Include handling for images
            for image in doc.inline_shapes:
                image_path = os.path.join(os.path.dirname(file_path), image._inline.graphic.graphicData.pic.nvPicPr.cNvPr.name)
                copyfile(image._inline.graphic.graphicData.pic.blipFill.blip.embed, os.path.join(os.path.dirname(file_path), os.path.basename(image_path)))
        elif file_ext == '.xlsx':
            wb = load_workbook(filename=file_path, read_only=True)
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                for row in sheet.iter_rows(values_only=True):
                    for cell in row:
                        if cell is not None:
                            text += str(cell) + '\t'  
                    text += '\n'
        # For further improvement currently have some issues works for few files not all'''
        # Commenting out the .pdf section temporarily
        # Uncomment below lines when you want to re-enable the .pdf section
        # elif file_ext == '.pdf':
        #     with open(file_path, 'rb') as file:
        #         pdf = PdfReader(file)
        #         for page_number in range(len(pdf.pages)):  # Updated line here
        #             text += pdf.pages[page_number].extract_text()  # And here
        elif file_ext == '.csv':
            with open(file_path, 'r', encoding='utf-8') as file:
                csv_reader = csv.reader(file)
                for row in csv_reader:
                    text += '\t'.join(row) + '\n'
        else:
            messagebox.showerror("Unsupported file type", "Please select a supported file type (.txt, .docx, .xlsx, .csv)")
            return None
    except Exception as e:
        logging.error(f"Error occurred while reading file: {e}")
        messagebox.showerror("Error", "An error occurred while reading the file.")
        return None
    return text


def shorten_filename(filename, max_length=50):
    # If filename is already shorter than the maximum length, return it as is
    if len(filename) <= max_length:
        return filename

    # Split the filename into parts separated by newline characters
    parts = filename.split('\n')

    # Take the first few words from each part to generate a shortened filename
    shortened_parts = [part.split()[:2] for part in parts if part.strip()]
    shortened_filename = ' '.join('_'.join(words) for words in shortened_parts)

    # Ensure the length of the shortened filename does not exceed the maximum length
    if len(shortened_filename) > max_length:
        shortened_filename = shortened_filename[:max_length]

    return shortened_filename


def save_file(file_path, translated_text, dest_lang):
    """
    Save translated text to a file.

    Parameters:
        file_path (str): The path to the original file.
        translated_text (str): The translated text to be saved.
        dest_lang (str): The destination language code.
    """
    try:
        dir_name = os.path.dirname(file_path)
        base_name = os.path.basename(file_path)
        name, ext = os.path.splitext(base_name)
        new_name = f"{name}_{dest_lang}{ext}"  # Append the destination language to the original file name
        new_path = os.path.join(dir_name, new_name)

        if ext == '.docx':
            doc = Document()
            for paragraph in translated_text.split('\n'):
                doc.add_paragraph(paragraph)
            # Include handling for images
            for image_file in os.listdir(dir_name):
                if image_file.endswith(('.png', '.jpg', '.jpeg', '.gif')):
                    doc.add_picture(os.path.join(dir_name, image_file))
            doc.save(new_path)
        elif ext == '.xlsx':
            # Create a new workbook
            wb_translated = Workbook()
            ws_translated = wb_translated.active

            # Write the translated text into the appropriate cells
            for row_idx, row in enumerate(translated_text.split('\n'), start=1):
                for col_idx, value in enumerate(row.split('\t'), start=1):
                    ws_translated.cell(row=row_idx, column=col_idx, value=value)

            # Save the workbook with the appropriate file extension
            wb_translated.save(new_path)

        # For further improvement currently have some issues works for few files not all
        # # Commenting out the .pdf section temporarily
        # Uncomment this line when you want to re-enable the .pdf section
        # elif ext == '.pdf':
        #     original_file = open(file_path, 'rb')
        #     original_pdf = PdfReader(original_file)

        #     # Create a new PDF file
        #     translated_pdf = PdfWriter()

        #     # Split translated text into lines
        #     lines = translated_text.split('\n')

        #     for i in range(len(original_pdf.pages)):
        #         # Create a new canvas for adding text
        #         packet = BytesIO()
        #         c = canvas.Canvas(packet, pagesize=letter)
        #         textobject = c.beginText()
        #         textobject.setTextOrigin(10,  730)
        #         textobject.setFont("Helvetica",  12)

        #         # Add each line of text to the canvas
        #         for line in lines:
        #             textobject.textLine(line)

        #         # Draw the text on the canvas
        #         c.drawText(textobject)
        #         c.save()

        #         # Move to the beginning of the StringIO buffer
        #         packet.seek(0)
        #         new_pdf = PdfReader(packet)

        #         # Get the original page
        #         # original_page = original_pdf.pages[i]

        #         # Merge the original page with the new text layer
        #         new_pdf.pages[0].add_transformation(Transformation().translate(0,  0))

        #         # Merge the original page with the new text layer
        #         # original_page.merge_page(new_pdf.pages[0], expand=True)

        #         # Add the merged page to the translated_pdf
        #         translated_pdf.add_page(new_pdf.pages[0])

        #     # Write the translated_pdf to a new file
        #     with open(new_path, 'wb') as output_file:
        #         translated_pdf.write(output_file)

        #     # Close the original file
        #     original_file.close()

        elif ext == '.csv':
            # Save the translated text as a CSV file
            with open(new_path, 'w', newline='', encoding='utf-8') as file:
                writer = csv.writer(file)
                for row in translated_text.split('\n'):
                    writer.writerow(row.split(','))

        else:
            # For other file types, save the translated text directly
            with open(new_path, 'w', encoding='utf-8') as file:
                file.write(translated_text)

        messagebox.showinfo("Success", "Translation saved successfully.")
    except Exception as e:
        print(e)
        logging.error(f"Error occurred while saving file: {e}")
        messagebox.showerror("Error", "An error occurred while saving the translated file.")


def translate_file():
    """
    Translate a file selected by the user.
    """
    try:
        file_path = filedialog.askopenfilename(filetypes=[("Text files", "*.txt"), ("Word files", "*.docx"), ("Excel files", "*.xlsx"), ("CSV files", "*.csv")])
        if not file_path:
            return

        # Read file
        original_text = read_file(file_path)
        if not original_text:
            return

        # Detect language
        detected_lang_code = detect_language(original_text)
        if not detected_lang_code:
            return
        detected_lang_name = LANGUAGES.get(detected_lang_code, "Unknown")
        messagebox.showinfo("Detected Language", f"The detected language is: {detected_lang_name}")

        # Fetch list of languages
        lang_dict = LANGUAGES

        # Choose language to translate from dropdown
        dest_lang_var = tk.StringVar(root)
        dest_lang_var.set("English")

        lang_menu = tk.OptionMenu(root, dest_lang_var, *lang_dict.values())
        lang_menu.place(relx=0.5, rely=0.5, anchor=tk.CENTER)

        def on_translate():
            selected_lang = dest_lang_var.get()
            translated_text = translate(original_text, selected_lang)
            if translated_text:
                save_file(file_path, translated_text, selected_lang)

        translate_button = tk.Button(root, text="Translate", command=on_translate)
        translate_button.place(relx=0.5, rely=0.7, anchor=tk.CENTER)
    except Exception as e:
        logging.error(f"Error occurred while translating file: {e}")
        messagebox.showerror("Error", "An error occurred while translating the file.")


def close_app():
    """
    Close the application.
    """
    root.destroy()


def main():
    """
    Main function to initialize the application window.
    """

    global root
    root = tk.Tk()
    root.title("File Translator")
    root.geometry("400x200")  # Set the size of the window

    canvas = tk.Canvas(root, bg="light blue", width=400, height=200)  # Background color
    canvas.pack()

    translate_button = tk.Button(root, text="Translate File", command=translate_file)
    translate_button.place(relx=0.5, rely=0.2, anchor=tk.CENTER)

    close_button = tk.Button(root, text="Close", command=close_app)
    close_button.place(relx=0.5, rely=0.9, anchor=tk.CENTER)

    root.mainloop()


if __name__ == "__main__":
    main()
